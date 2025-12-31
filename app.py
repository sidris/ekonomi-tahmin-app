import streamlit as st
from supabase import create_client, Client
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from fpdf import FPDF
from fpdf.fonts import FontFace
import tempfile
import os
import io
import datetime
import time
import requests
import xlsxwriter
import smtplib
from email.mime.text import MIMEText
import random

# --- 1. AYARLAR VE TASARIM ---
st.set_page_config(page_title="Finansal Tahmin Terminali", layout="wide", page_icon="📊", initial_sidebar_state="expanded")

# --- CUSTOM CSS ---
st.markdown("""
<style>
    .stMetric { background-color: #ffffff; border: 1px solid #e0e0e0; padding: 15px; border-radius: 10px; box-shadow: 0 2px 4px rgba(0,0,0,0.05); } 
    .stButton button { width: 100%; border-radius: 8px; font-weight: 600; } 
    div[data-testid="stExpander"] { border: 1px solid #e0e0e0; border-radius: 8px; background-color: white; } 
    div[data-testid="stDataFrame"] { width: 100%; }
    
    /* Login Ekranı */
    .login-container {
        background-color: white;
        padding: 3rem;
        border-radius: 15px;
        box-shadow: 0 10px 25px rgba(0,0,0,0.1);
        text-align: center;
        margin-top: 50px;
    }
    .login-header {
        color: #1E3A8A;
        font-family: 'Helvetica Neue', sans-serif;
        font-weight: 700;
        margin-bottom: 1.5rem;
    }
</style>
""", unsafe_allow_html=True)

# --- KÜTÜPHANE KONTROLÜ ---
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.error("Lütfen gerekli kütüphaneleri yükleyin: pip install python-docx xlsxwriter requests fpdf plotly pandas supabase openpyxl")
    st.stop()

# --- BAĞLANTI VE AYARLAR ---
try:
    url = st.secrets["SUPABASE_URL"]
    key = st.secrets["SUPABASE_KEY"]
    SITE_SIFRESI = st.secrets["APP_PASSWORD"]
    EVDS_API_KEY = st.secrets.get("EVDS_KEY", None)
    
    # E-posta Ayarları
    SMTP_EMAIL = st.secrets.get("SMTP_EMAIL", None)
    SMTP_PASSWORD = st.secrets.get("SMTP_PASSWORD", None)
    
    supabase: Client = create_client(url, key)
except Exception as e:
    st.error(f"Lütfen secrets ayarlarını kontrol edin: {e}")
    st.stop()

TABLE_TAHMIN = "tahminler4"
TABLE_KATILIMCI = "katilimcilar"
EVDS_BASE = "https://evds2.tcmb.gov.tr/service/evds"
EVDS_TUFE_SERIES = "TP.FG.J0" 

# --- YARDIMCI FONKSİYONLAR ---

def get_period_list():
    years = range(2024, 2033)
    months = ["01", "02", "03", "04", "05", "06", "07", "08", "09", "10", "11", "12"]
    period_list = []
    for y in years:
        for m in months:
            period_list.append(f"{y}-{m}")
    return period_list

tum_donemler = get_period_list()

def normalize_name(name): return name.strip().title() if name else ""

def safe_int(val):
    try: return int(float(val)) if pd.notnull(val) else 0
    except: return 0

def clean_and_sort_data(df):
    if df.empty: return df
    numeric_cols = ["tahmin_ppk_faiz", "min_ppk_faiz", "max_ppk_faiz", "tahmin_yilsonu_faiz", 
                    "min_yilsonu_faiz", "max_yilsonu_faiz", "tahmin_aylik_enf", "min_aylik_enf", 
                    "max_aylik_enf", "tahmin_yillik_enf", "min_yillik_enf", "max_yillik_enf", 
                    "tahmin_yilsonu_enf", "min_yilsonu_enf", "max_yilsonu_enf", "katilimci_sayisi"]
    for col in numeric_cols:
        if col in df.columns: df[col] = pd.to_numeric(df[col], errors='coerce')
    if "donem" in df.columns:
        df["donem_date"] = pd.to_datetime(df["donem"], format="%Y-%m", errors='coerce')
        df = df.sort_values(by="donem_date")
    if "tahmin_tarihi" in df.columns:
        df["tahmin_tarihi"] = pd.to_datetime(df["tahmin_tarihi"])
    return df

def parse_range_input(text_input, default_median=0.0):
    if not text_input or text_input.strip() == "": return default_median, 0.0, 0.0, False
    try:
        text = text_input.replace(',', '.')
        parts = []
        if '-' in text: parts = text.split('-')
        elif '/' in text: parts = text.split('/')
        if len(parts) == 2:
            v1, v2 = float(parts[0].strip()), float(parts[1].strip())
            return (v1+v2)/2, min(v1, v2), max(v1, v2), True
    except: pass
    return default_median, 0.0, 0.0, False

def upsert_tahmin(user, period, category, forecast_date, link, data_dict):
    if isinstance(forecast_date, str):
        date_str = forecast_date
    else:
        date_str = forecast_date.strftime("%Y-%m-%d")
    
    check_res = supabase.table(TABLE_TAHMIN).select("*").eq("kullanici_adi", user).eq("donem", period).execute()
    
    existing_data = {}
    record_id = None
    
    if check_res.data:
        existing_data = check_res.data[0]
        record_id = existing_data['id']
        for k in ['id', 'created_at', 'kullanici_adi', 'donem']: 
            if k in existing_data: del existing_data[k]

    new_input_data = {k: v for k, v in data_dict.items() if v is not None and v != 0 and v != ""}
    final_data = existing_data.copy()
    final_data.update(new_input_data)
    
    final_data.update({
        "kullanici_adi": user, 
        "donem": period, 
        "kategori": category, 
        "tahmin_tarihi": date_str
    })
    
    if link: final_data["kaynak_link"] = link

    if record_id:
        supabase.table(TABLE_TAHMIN).update(final_data).eq("id", record_id).execute()
        return "updated"
    else:
        supabase.table(TABLE_TAHMIN).insert(final_data).execute()
        return "inserted"

def to_excel(df):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer: df.to_excel(writer, index=False, sheet_name='Tahminler')
    return output.getvalue()

# --- E-POSTA GÖNDERİM FONKSİYONU ---
def send_verification_email(code):
    if not SMTP_EMAIL or not SMTP_PASSWORD:
        return False, "SMTP Ayarları eksik! Lütfen secrets.toml dosyasını kontrol edin."
    
    try:
        msg = MIMEText(f"""
        Merhaba,
        
        Finansal Tahmin Terminali veritabanını TAMAMEN SİLMEK için bir talep oluşturuldu.
        
        Onay Kodunuz: {code}
        
        Eğer bu işlemi siz başlatmadıysanız, lütfen bu mesajı dikkate almayın.
        """)
        msg['Subject'] = '🚨 VERİ SİLME ONAY KODU'
        msg['From'] = SMTP_EMAIL
        msg['To'] = "s.idrisoglu@gmail.com" # Sabitlenen mail adresi

        # Gmail SSL portu 465
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as server:
            server.login(SMTP_EMAIL, SMTP_PASSWORD)
            server.sendmail(SMTP_EMAIL, "s.idrisoglu@gmail.com", msg.as_string())
        return True, "Kod gönderildi."
    except Exception as e:
        return False, f"Hata: {str(e)}"

# =========================================================
# VERİ ÇEKME MOTORU
# =========================================================
def _evds_headers(api_key: str) -> dict:
    return {"key": api_key, "User-Agent": "Mozilla/5.0"}

def _evds_url_single(series_code: str, start_date: datetime.date, end_date: datetime.date, formulas: int | None) -> str:
    s = start_date.strftime("%d-%m-%Y")
    e = end_date.strftime("%d-%m-%Y")
    url = f"{EVDS_BASE}/series={series_code}&startDate={s}&endDate={e}&type=json"
    if formulas is not None: url += f"&formulas={int(formulas)}"
    return url

@st.cache_data(ttl=600)
def fetch_evds_tufe_monthly_yearly(api_key: str, start_date: datetime.date, end_date: datetime.date) -> tuple[pd.DataFrame, str | None]:
    if not api_key: return pd.DataFrame(), "EVDS_KEY eksik."
    try:
        results = {}
        for formulas, out_col in [(1, "TUFE_Aylik"), (3, "TUFE_Yillik")]:
            url = _evds_url_single(EVDS_TUFE_SERIES, start_date, end_date, formulas=formulas)
            r = requests.get(url, headers=_evds_headers(api_key), timeout=25)
            if r.status_code != 200: continue
            js = r.json()
            items = js.get("items", [])
            if not items: continue
            df = pd.DataFrame(items)
            if "Tarih" not in df.columns: continue
            df["Tarih_dt"] = pd.to_datetime(df["Tarih"], dayfirst=True, errors="coerce")
            if df["Tarih_dt"].isnull().all(): df["Tarih_dt"] = pd.to_datetime(df["Tarih"], format="%Y-%m", errors="coerce")
            df = df.dropna(subset=["Tarih_dt"]).sort_values("Tarih_dt")
            df["Donem"] = df["Tarih_dt"].dt.strftime("%Y-%m")
            val_cols = [c for c in df.columns if c not in ["Tarih", "UNIXTIME", "Tarih_dt", "Donem"]]
            if not val_cols: continue
            results[out_col] = pd.DataFrame({"Tarih": df["Tarih_dt"].dt.strftime("%d-%m-%Y"), "Donem": df["Donem"], out_col: pd.to_numeric(df[val_cols[0]], errors="coerce")})

        df_m = results.get("TUFE_Aylik", pd.DataFrame())
        df_y = results.get("TUFE_Yillik", pd.DataFrame())
        if df_m.empty and df_y.empty: return pd.DataFrame(), "Veri bulunamadı."
        if df_m.empty: out = df_y
        elif df_y.empty: out = df_m
        else: out = pd.merge(df_m, df_y, on=["Tarih", "Donem"], how="outer")
        return out.sort_values(["Donem", "Tarih"]), None
    except Exception as e: return pd.DataFrame(), str(e)

@st.cache_data(ttl=600)
def fetch_bis_cbpol_tr(start_date: datetime.date, end_date: datetime.date) -> tuple[pd.DataFrame, str | None]:
    try:
        s = start_date.strftime("%Y-%m-%d"); e = end_date.strftime("%Y-%m-%d")
        url = f"https://stats.bis.org/api/v1/data/WS_CBPOL/D.TR?format=csv&startPeriod={s}&endPeriod={e}"
        r = requests.get(url, timeout=25)
        if r.status_code >= 400: return pd.DataFrame(), f"BIS HTTP {r.status_code}"
        df = pd.read_csv(io.StringIO(r.content.decode("utf-8", errors="ignore")))
        df.columns = [c.strip().upper() for c in df.columns]
        out = df[["TIME_PERIOD", "OBS_VALUE"]].copy()
        out["TIME_PERIOD"] = pd.to_datetime(out["TIME_PERIOD"], errors="coerce")
        out = out.dropna(subset=["TIME_PERIOD"])
        out["Donem"] = out["TIME_PERIOD"].dt.strftime("%Y-%m")
        out["Tarih"] = out["TIME_PERIOD"].dt.strftime("%d-%m-%Y")
        out["REPO_RATE"] = pd.to_numeric(out["OBS_VALUE"], errors="coerce")
        return out[["Donem", "REPO_RATE"]].sort_values(["Donem"]), None
    except Exception as e: return pd.DataFrame(), str(e)

def fetch_market_data_adapter(api_key, start_date, end_date):
    df_inf, err1 = fetch_evds_tufe_monthly_yearly(api_key, start_date, end_date)
    df_pol, err2 = fetch_bis_cbpol_tr(start_date, end_date)
    if df_inf.empty and df_pol.empty: return pd.DataFrame(), f"Veri Yok: {err1} | {err2}"
    combined = pd.DataFrame()
    if not df_inf.empty and not df_pol.empty:
        df_pol_monthly = df_pol.groupby("Donem").last().reset_index()[['Donem', 'REPO_RATE']]
        combined = pd.merge(df_inf, df_pol_monthly, on="Donem", how="outer")
    elif not df_inf.empty: combined = df_inf; combined['REPO_RATE'] = None
    elif not df_pol.empty: combined = df_pol.rename(columns={'REPO_RATE': 'REPO_RATE'}); combined['TUFE_Aylik'] = None; combined['TUFE_Yillik'] = None
    combined = combined.rename(columns={'REPO_RATE': 'PPK Faizi', 'TUFE_Aylik': 'Aylık TÜFE', 'TUFE_Yillik': 'Yıllık TÜFE'})
    if 'Tarih' not in combined.columns and 'Donem' in combined.columns: combined['Tarih'] = combined['Donem'] + "-01"
    return combined, None

# --- PDF & EXCEL EXPORT ARAÇLARI ---
def check_and_download_font():
    try:
        if not os.path.exists("DejaVuSans.ttf"):
            r = requests.get("https://github.com/google/fonts/raw/main/ofl/dejavusans/DejaVuSans-Regular.ttf", timeout=10)
            with open("DejaVuSans.ttf", 'wb') as f: f.write(r.content)
        if not os.path.exists("DejaVuSans-Bold.ttf"):
            r = requests.get("https://github.com/google/fonts/raw/main/ofl/dejavusans/DejaVuSans-Bold.ttf", timeout=10)
            with open("DejaVuSans-Bold.ttf", 'wb') as f: f.write(r.content)
        return "DejaVuSans.ttf", "DejaVuSans-Bold.ttf"
    except: return None, None

def safe_str(text, fallback):
    if not isinstance(text, str): return str(text)
    if fallback:
        tr = {'ğ':'g','Ğ':'G','ş':'s','Ş':'S','ı':'i','İ':'I','ö':'o','Ö':'O','ü':'u','Ü':'U','ç':'c','Ç':'C'}
        for k,v in tr.items(): text = text.replace(k,v)
    return text

def create_custom_pdf_report(report_data):
    fr, fb = check_and_download_font()
    use_cust = (fr is not None); font = "DejaVu" if use_cust else "Helvetica"; fallback = not use_cust
    class RPT(FPDF):
        def header(self):
            self.ln(25)
        def footer(self):
            self.set_y(-15); self.set_font(font, '', 8); self.set_text_color(128); self.cell(0, 10, f'Sayfa {self.page_no()}', align='C')
    pdf = RPT(); 
    if use_cust: pdf.add_font("DejaVu", "", fr, uni=True); pdf.add_font("DejaVu", "B", fb, uni=True)
    pdf.add_page(); pdf.set_text_color(0)
    pdf.set_font(font, 'B', 20); pdf.cell(0, 10, safe_str(report_data['title'], fallback), ln=True)
    pdf.set_font(font, '', 12); pdf.set_text_color(80); pdf.cell(0, 8, safe_str(report_data['unit'], fallback), ln=True)
    pdf.set_text_color(0); pdf.set_font(font, '', 10); pdf.cell(0, 8, safe_str(report_data['date'], fallback), ln=True, align='R'); pdf.ln(5)
    if report_data['body']: pdf.set_font(font, '', 11); pdf.multi_cell(0, 6, safe_str(report_data['body'], fallback)); pdf.ln(10)
    for block in report_data['content_blocks']:
        if pdf.get_y() > 240: pdf.add_page()
        if block.get('title'): pdf.set_font(font, 'B', 12); pdf.set_text_color(200, 0, 0); pdf.cell(0, 10, safe_str(block['title'], fallback), ln=True); pdf.set_text_color(0); pdf.ln(2)
        if block['type'] == 'chart':
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as t:
                try: block['fig'].write_image(t.name, width=1000, height=500, scale=2); pdf.image(t.name, x=15, w=180); pdf.ln(5)
                except: pass
            try: os.remove(t.name)
            except: pass
        elif block['type'] == 'table':
            df = block['df']; pdf.set_font(font, '', 8)
            with pdf.table() as tbl:
                r = tbl.row()
                for c in df.columns: r.cell(safe_str(str(c), fallback), style=FontFace(emphasis="BOLD", color=255, fill_color=(200, 50, 50)))
                for _, dr in df.iterrows():
                    r = tbl.row(); 
                    for item in dr: r.cell(safe_str(str(item), fallback))
            pdf.ln(10)
    return bytes(pdf.output())

def create_word_report(report_data):
    doc = Document()
    title = doc.add_heading(report_data['title'], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_unit = p_info.add_run(report_data['unit'] + "\n")
    run_unit.bold = True; run_unit.font.size = Pt(12)
    run_date = p_info.add_run(report_data['date'])
    run_date.italic = True
    doc.add_paragraph("")
    if report_data['body']: p_body = doc.add_paragraph(report_data['body']); p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for block in report_data['content_blocks']:
        doc.add_paragraph("")
        if block.get('title'): h = doc.add_heading(block['title'], level=2); h.runs[0].font.color.rgb = RGBColor(180, 0, 0)
        if block['type'] == 'chart':
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmpfile:
                try: block['fig'].write_image(tmpfile.name, width=1000, height=500, scale=2); doc.add_picture(tmpfile.name, width=Inches(6.5))
                except: pass
            try: os.remove(tmpfile.name)
            except: pass
        elif block['type'] == 'table':
            df_table = block['df']; table = doc.add_table(rows=1, cols=len(df_table.columns)); table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(df_table.columns): hdr_cells[i].text = str(col_name)
            for _, row in df_table.iterrows():
                row_cells = table.add_row().cells
                for i, item in enumerate(row): row_cells[i].text = str(item)
    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()

def create_excel_dashboard(df_source):
    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {'in_memory': True})
    bold = workbook.add_format({'bold': 1}); date_fmt = workbook.add_format({'num_format': 'dd/mm/yyyy'})
    ws_raw = workbook.add_worksheet("Ham Veri"); ws_raw.write_row('A1', df_source.columns, bold)
    for r, row in enumerate(df_source.values):
        for c, val in enumerate(row):
            if pd.isna(val): ws_raw.write_string(r+1, c, "")
            elif isinstance(val, (datetime.date, datetime.datetime)): ws_raw.write_datetime(r+1, c, val, date_fmt)
            else: ws_raw.write(r+1, c, val)
    def create_chart(metric, name, title):
        try:
            piv = df_source.sort_values("donem_date").pivot(index='donem', columns='gorunen_isim', values=metric)
            ws = workbook.add_worksheet(name); ws.write('A1', 'Dönem', bold); ws.write_row('B1', piv.columns, bold); ws.write_column('A2', piv.index)
            for i, col in enumerate(piv.columns):
                for r_idx, val in enumerate(piv[col]):
                    if pd.notnull(val): ws.write_number(r_idx+1, i+1, val)
            chart = workbook.add_chart({'type': 'line'})
            for i in range(len(piv.columns)):
                chart.add_series({'name': [name, 0, i+1], 'categories': [name, 1, 0, len(piv), 0], 'values': [name, 1, i+1, len(piv), i+1]})
            chart.set_title({'name': title}); ws.insert_chart('E2', chart)
        except: pass
    create_chart('tahmin_ppk_faiz', 'PPK Graf', 'PPK Beklentileri')
    create_chart('tahmin_yilsonu_enf', 'Enf Graf', 'Yıl Sonu Enflasyon')
    workbook.close()
    return output.getvalue()

# --- LOGIN ---
if 'giris_yapildi' not in st.session_state: st.session_state['giris_yapildi'] = False
if not st.session_state['giris_yapildi']:
    c1, c2, c3 = st.columns([1, 2, 1])
    with c2:
        st.markdown("""<div class="login-container"><h1 class="login-header">📊 Finansal Tahmin Terminali</h1><p style="color: #666; margin-bottom: 20px;">Lütfen erişim için şifrenizi giriniz.</p></div>""", unsafe_allow_html=True)
        with st.form("login_form"):
            pwd = st.text_input("Şifre", type="password")
            st.markdown("<br>", unsafe_allow_html=True)
            if st.form_submit_button("Giriş Yap", type="primary"):
                if pwd == SITE_SIFRESI: st.session_state['giris_yapildi'] = True; st.rerun()
                else: st.error("Hatalı Şifre!")
    st.stop()

# --- SIDEBAR ---
with st.sidebar:
    st.title("📊 Menü")
    page = st.radio("Git:", ["Gelişmiş Veri Havuzu (Yönetim)", "Dashboard", "🔥 Isı Haritası", "📈 Piyasa Verileri (EVDS)", "📄 Rapor Oluştur", "📥 Toplu Veri Yükleme (Excel)", "PPK Girişi", "Enflasyon Girişi", "Katılımcı Yönetimi"])

def get_participant_selection():
    res = supabase.table(TABLE_KATILIMCI).select("*").order("ad_soyad").execute()
    df = pd.DataFrame(res.data)
    if df.empty: st.error("Lütfen önce Katılımcı ekleyin."); return None, None, None
    df['disp'] = df.apply(lambda x: f"{x['ad_soyad']} ({x['anket_kaynagi']})" if x['anket_kaynagi'] else x['ad_soyad'], axis=1)
    name_map = dict(zip(df['disp'], df['ad_soyad']))
    sel = st.selectbox("Katılımcı Seç", df["disp"].unique())
    row = df[df["ad_soyad"] == name_map[sel]].iloc[0]
    return name_map[sel], row['kategori'], sel

# ========================================================
# SAYFA: GELİŞMİŞ VERİ HAVUZU (YÖNETİM)
# ========================================================
if page == "Gelişmiş Veri Havuzu (Yönetim)":
    st.title("🗃️ Veri Havuzu ve Yönetim Paneli")
    res_t = supabase.table(TABLE_TAHMIN).select("*").order("tahmin_tarihi", desc=True).limit(2000).execute()
    df_t = pd.DataFrame(res_t.data)
    if not df_t.empty:
        df_t = clean_and_sort_data(df_t)
        res_k = supabase.table(TABLE_KATILIMCI).select("ad_soyad", "kategori", "anket_kaynagi").execute()
        df_k = pd.DataFrame(res_k.data)
        if not df_k.empty:
            df_full = pd.merge(df_t, df_k, left_on="kullanici_adi", right_on="ad_soyad", how="left")
            df_full['kategori'] = df_full['kategori_y'].fillna('Bireysel')
            df_full['anket_kaynagi'] = df_full['anket_kaynagi'].fillna('-')
            df_full['tahmin_tarihi'] = pd.to_datetime(df_full['tahmin_tarihi'])
            with st.container():
                c1, c2, c3, c4, c5 = st.columns(5)
                sel_cat = c1.selectbox("Kategori", ["Tümü"] + list(df_full['kategori'].unique()))
                sel_period = c2.selectbox("Dönem", ["Tümü"] + sorted(list(df_full['donem'].unique()), reverse=True))
                sel_user = c3.selectbox("Katılımcı", ["Tümü"] + sorted(list(df_full['kullanici_adi'].unique())))
                sort_option = c4.selectbox("Sıralama", ["Tarih (Yeniden Eskiye)", "Tarih (Eskiden Yeniye)", "Son Eklenen (ID)"])
                admin_mode = c5.toggle("🛠️ Yönetici Modu")
            df_f = df_full.copy()
            if sel_cat != "Tümü": df_f = df_f[df_f['kategori'] == sel_cat]
            if sel_period != "Tümü": df_f = df_f[df_f['donem'] == sel_period]
            if sel_user != "Tümü": df_f = df_f[df_f['kullanici_adi'] == sel_user]
            if sort_option == "Tarih (Yeniden Eskiye)": df_f = df_f.sort_values(by="tahmin_tarihi", ascending=False)
            elif sort_option == "Tarih (Eskiden Yeniye)": df_f = df_f.sort_values(by="tahmin_tarihi", ascending=True)
            else: df_f = df_f.sort_values(by="id", ascending=False)
            if not admin_mode:
                st.markdown("---")
                cols = ["tahmin_tarihi", "donem", "kullanici_adi", "kategori", "anket_kaynagi", "kaynak_link", "katilimci_sayisi", "tahmin_ppk_faiz", "tahmin_yilsonu_faiz", "tahmin_aylik_enf", "tahmin_yillik_enf", "tahmin_yilsonu_enf"]
                final_cols = [c for c in cols if c in df_f.columns]
                col_cfg = {"kaynak_link": st.column_config.LinkColumn("Link", display_text="🔗"), "tahmin_tarihi": st.column_config.DateColumn("Tarih", format="DD.MM.YYYY"), **{c: st.column_config.NumberColumn(c, format="%.2f") for c in final_cols if "tahmin" in c}}
                st.dataframe(df_f[final_cols].sort_values(by="tahmin_tarihi", ascending=False), column_config=col_cfg, use_container_width=True, height=600)
                if not df_f.empty:
                    df_ex = df_f.copy(); df_ex['tahmin_tarihi'] = df_ex['tahmin_tarihi'].dt.strftime('%Y-%m-%d')
                    st.download_button("📥 Excel İndir", to_excel(df_ex), f"Veri_{sel_user}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
            else:
                if 'admin_ok' not in st.session_state: st.session_state['admin_ok'] = False
                if not st.session_state['admin_ok']:
                    with st.form("admin_login_form"):
                        st.subheader("Yönetici Girişi")
                        if st.form_submit_button("Giriş Yap") and st.text_input("Şifre", type="password") == "Admin": st.session_state['admin_ok'] = True; st.rerun()
                else:
                    if 'edit_target' in st.session_state:
                        t = st.session_state['edit_target']
                        with st.form("full_edit_form"):
                            st.subheader(f"Düzenle: {t['kullanici_adi']} ({t['donem']})"); st.warning("⚠️ 'Ana Tahmin' bölümü mevcut kaydın üzerine yazar. 'İleri Vadeli' bölümü yeni kayıt oluşturur.")
                            c1, c2, c3 = st.columns(3); nd = c1.date_input("Tarih", pd.to_datetime(t.get('tahmin_tarihi')).date()); ndo = c2.selectbox("Dönem", tum_donemler, index=tum_donemler.index(t['donem']) if t['donem'] in tum_donemler else 0); nl = c3.text_input("Link", t.get('kaynak_link') or "")
                            def g(k): return float(t.get(k) or 0)
                            original_n = safe_int(t.get('katilimci_sayisi'))
                            tp, te = st.tabs(["Faiz", "Enflasyon"])
                            with tp:
                                c1, c2 = st.columns(2); r_ppk = c1.text_input("Aralık PPK", key="r_ppk"); v_ppk = c1.number_input("Medyan PPK", value=g('tahmin_ppk_faiz'), step=0.25); r_yf = c2.text_input("Aralık YS Faiz", key="r_yf"); v_yf = c2.number_input("Medyan YS Faiz", value=g('tahmin_yilsonu_faiz'), step=0.25)
                                with st.expander("Detaylar", expanded=True):
                                    ec1, ec2, ec3 = st.columns(3); mn_ppk = ec1.number_input("Min PPK", value=g('min_ppk_faiz'), step=0.25); mx_ppk = ec1.number_input("Max PPK", value=g('max_ppk_faiz'), step=0.25); mn_yf = ec2.number_input("Min YS Faiz", value=g('min_yilsonu_faiz'), step=0.25); mx_yf = ec2.number_input("Max YS Faiz", value=g('max_yilsonu_faiz'), step=0.25); nk_faiz = ec3.number_input("N", value=original_n, step=1, key="nk_edit_faiz")
                                md, mn, mx, ok = parse_range_input(r_ppk, v_ppk); 
                                if ok: v_ppk, mn_ppk, mx_ppk = md, mn, mx
                                md2, mn2, mx2, ok2 = parse_range_input(r_yf, v_yf)
                                if ok2: v_yf, mn_yf, mx_yf = md2, mn2, mx2
                            with te:
                                c1, c2, c3 = st.columns(3); r_ay = c1.text_input("Aralık Ay", key="r_ay"); v_ay = c1.number_input("Medyan Ay", value=g('tahmin_aylik_enf'), step=0.01); r_yil = c2.text_input("Aralık Yıllık", key="r_yil"); v_yil = c2.number_input("Medyan Yıllık", value=g('tahmin_yillik_enf'), step=0.01); r_ys = c3.text_input("Aralık YS", key="r_ys"); v_ys = c3.number_input("Medyan YS", value=g('tahmin_yilsonu_enf'), step=0.01)
                                with st.expander("Detaylar", expanded=True):
                                    ec1, ec2, ec3 = st.columns(3); mn_ay = ec1.number_input("Min Ay", value=g('min_aylik_enf'), step=0.01); mx_ay = ec1.number_input("Max Ay", value=g('max_aylik_enf'), step=0.01); mn_yil = ec2.number_input("Min Yıllık", value=g('min_yillik_enf'), step=0.01); mx_yil = ec2.number_input("Max Yıllık", value=g('max_yillik_enf'), step=0.01); mn_ys = ec3.number_input("Min YS", value=g('min_yilsonu_enf'), step=0.01); mx_ys = ec3.number_input("Max YS", value=g('max_yilsonu_enf'), step=0.01); nk_enf = st.number_input("N - Enflasyon", value=original_n, step=1, key="nk_edit_enf")
                                md1, mn1, mx1, ok1 = parse_range_input(r_ay, v_ay); 
                                if ok1: v_ay, mn_ay, mx_ay = md1, mn1, mx1
                                md2, mn2, mx2, ok2 = parse_range_input(r_yil, v_yil)
                                if ok2: v_yil, mn_yil, mx_yil = md2, mn2, mx2
                                md3, mn3, mx3, ok3 = parse_range_input(r_ys, v_ys)
                                if ok3: v_ys, mn_ys, mx_ys = md3, mn3, mx3
                            
                            st.markdown("---"); st.markdown("#### 📅 İleri Vadeli Ekle (Opsiyonel)"); fe1, fe2, fe3 = st.columns(3); future_donem = fe1.selectbox("Hedef", tum_donemler, index=0); future_val_enf = fe2.number_input("Gelecek Enflasyon", step=0.01); future_val_ppk = fe3.number_input("Gelecek PPK", step=0.25)
                            
                            if st.form_submit_button("💾 Kaydet"):
                                def cv(v): 
                                    try: val = float(v); return val if (pd.notnull(val) and val != 0) else None
                                    except: return None
                                final_nk = nk_enf if nk_enf != original_n else nk_faiz
                                upd = {"tahmin_tarihi": nd.strftime('%Y-%m-%d'), "donem": ndo, "kaynak_link": nl if nl else None, "katilimci_sayisi": int(final_nk), "tahmin_ppk_faiz": cv(v_ppk), "min_ppk_faiz": cv(mn_ppk), "max_ppk_faiz": cv(mx_ppk), "tahmin_yilsonu_faiz": cv(v_yf), "min_yilsonu_faiz": cv(mn_yf), "max_yilsonu_faiz": cv(mx_yf), "tahmin_aylik_enf": cv(v_ay), "min_aylik_enf": cv(mn_ay), "max_aylik_enf": cv(mx_ay), "tahmin_yillik_enf": cv(v_yil), "min_yillik_enf": cv(mn_yil), "max_yillik_enf": cv(mx_yil), "tahmin_yilsonu_enf": cv(v_ys), "min_yilsonu_enf": cv(mn_ys), "max_yilsonu_enf": cv(mx_ys)}
                                supabase.table(TABLE_TAHMIN).update(upd).eq("id", int(t['id'])).execute()
                                if future_val_enf > 0 or future_val_ppk > 0:
                                    fd = {"katilimci_sayisi": int(final_nk)}
                                    if future_val_enf > 0: fd.update({"tahmin_yilsonu_enf": future_val_enf, "tahmin_yillik_enf": future_val_enf})
                                    if future_val_ppk > 0: fd.update({"tahmin_ppk_faiz": future_val_ppk, "tahmin_yilsonu_faiz": future_val_ppk})
                                    upsert_tahmin(t['kullanici_adi'], future_donem, t['kategori'] or 'Bireysel', nd, nl, fd)
                                st.success("Güncellendi!"); time.sleep(1); del st.session_state['edit_target']; st.rerun()
                        if st.button("İptal"): del st.session_state['edit_target']; st.rerun()
                    else:
                        st.markdown("---"); h1, h2, h3, h4 = st.columns([2, 4, 1, 1]); h1.caption("Tarih"); h2.caption("Katılımcı / Dönem")
                        for idx, row in df_f.iterrows():
                            with st.container():
                                c1, c2, c3, c4 = st.columns([2, 4, 1, 1]); c1.write(row['tahmin_tarihi'].strftime('%d.%m.%Y')); c2.markdown(f"**{row['kullanici_adi']}** | {row['donem']}")
                                if c3.button("✏️", key=f"e{row['id']}"): st.session_state['edit_target'] = row; st.rerun()
                                if c4.button("🗑️", key=f"d{row['id']}"): supabase.table(TABLE_TAHMIN).delete().eq("id", int(row['id'])).execute(); st.rerun()

# ========================================================
# SAYFA: DASHBOARD
# ========================================================
elif page == "Dashboard":
    st.header("Piyasa Analiz Dashboardu")
    res_t = supabase.table(TABLE_TAHMIN).select("*").order("tahmin_tarihi", desc=True).limit(2500).execute()
    df_t = pd.DataFrame(res_t.data)
    res_k = supabase.table(TABLE_KATILIMCI).select("ad_soyad", "anket_kaynagi").execute()
    df_k = pd.DataFrame(res_k.data)

    if not df_t.empty and not df_k.empty:
        df_t = clean_and_sort_data(df_t)
        df_t['tahmin_tarihi'] = pd.to_datetime(df_t['tahmin_tarihi'])
        df_t = df_t.sort_values(by='tahmin_tarihi')
        
        dash_evds_start = datetime.date(2023, 1, 1); dash_evds_end = datetime.date(2025, 12, 31)
        realized_df, err = fetch_market_data_adapter(EVDS_API_KEY, dash_evds_start, dash_evds_end)
        
        realized_dict = {}
        if not realized_df.empty:
            for _, row in realized_df.iterrows():
                realized_dict[row['Donem']] = {'ppk': row.get('PPK Faizi'), 'enf_ay': row.get('Aylık TÜFE'), 'enf_yil': row.get('Yıllık TÜFE')}

        df_history = pd.merge(df_t, df_k, left_on="kullanici_adi", right_on="ad_soyad", how="inner")
        df_latest_raw = df_t.drop_duplicates(subset=['kullanici_adi', 'donem'], keep='last')
        df_latest = pd.merge(df_latest_raw, df_k, left_on="kullanici_adi", right_on="ad_soyad", how="inner")
        
        for d in [df_history, df_latest]:
            d['gorunen_isim'] = d.apply(lambda x: f"{x['kullanici_adi']} ({x['anket_kaynagi']})" if pd.notnull(x['anket_kaynagi']) and x['anket_kaynagi'] != '' else x['kullanici_adi'], axis=1)
            d['hover_text'] = d.apply(lambda x: f"Tarih: {x['tahmin_tarihi'].strftime('%d-%m-%Y')}<br>N={int(x['katilimci_sayisi'])}" if pd.notnull(x['katilimci_sayisi']) else "", axis=1)
            d['kategori'] = d['kategori'].fillna('Bireysel'); d['anket_kaynagi'] = d['anket_kaynagi'].fillna('-'); d['yil'] = d['donem'].apply(lambda x: x.split('-')[0])

        c1, c2, c3 = st.columns(3)
        c1.metric("Toplam Katılımcı", df_latest['kullanici_adi'].nunique())
        c2.metric("Toplam Tahmin Verisi", len(df_latest))
        c3.metric("Son Güncelleme", df_latest['tahmin_tarihi'].max().strftime('%d.%m.%Y'))
        st.markdown("---")

        st.subheader("🏆 Dönemin En İsabetli Tahmincileri")
        if not realized_df.empty:
            available_realized_periods = sorted(realized_df['Donem'].unique().tolist(), reverse=True)
            with st.expander("⚙️ Performans Analizi Ayarları (Tarih Aralığı)", expanded=True):
                col_p1, col_p2 = st.columns(2)
                p_end = col_p1.selectbox("Bitiş Dönemi", available_realized_periods, index=0)
                remain_periods = [p for p in available_realized_periods if p <= p_end]
                p_start = col_p2.selectbox("Başlangıç Dönemi", remain_periods, index=min(2, len(remain_periods)-1))
            
            mask_real = (realized_df['Donem'] >= p_start) & (realized_df['Donem'] <= p_end)
            target_real_df = realized_df[mask_real].copy()
            
            if not target_real_df.empty:
                perf_df = pd.merge(df_latest, target_real_df, left_on="donem", right_on="Donem", how="inner")
                perf_df['err_ppk'] = (perf_df['tahmin_ppk_faiz'] - perf_df['PPK Faizi']).abs()
                perf_df['err_enf_ay'] = (perf_df['tahmin_aylik_enf'] - perf_df['Aylık TÜFE']).abs()
                if 'tahmin_yillik_enf' in perf_df.columns: perf_df['val_enf_yil'] = perf_df['tahmin_yillik_enf'].fillna(perf_df['tahmin_yilsonu_enf'])
                else: perf_df['val_enf_yil'] = perf_df['tahmin_yilsonu_enf']
                perf_df['err_enf_yil'] = (perf_df['val_enf_yil'] - perf_df['Yıllık TÜFE']).abs()

                c_best1, c_best2, c_best3 = st.columns(3)
                def show_champion_card(col_obj, title, err_col, unit, icon, pred_col, act_col):
                    valid_df = perf_df.dropna(subset=[err_col])
                    if valid_df.empty: col_obj.warning(f"{title}\nVeri yok."); return
                    leaderboard = valid_df.groupby('gorunen_isim').agg({err_col: 'mean', pred_col: 'mean', act_col: 'mean', 'donem': 'count'}).reset_index()
                    leaderboard = leaderboard.sort_values(by=[err_col, 'donem'], ascending=[True, False])
                    winner = leaderboard.iloc[0]
                    col_obj.success(f"{icon} **{title}**\n\n🥇 **{winner['gorunen_isim']}**\n\nOrt. Tahmin: **%{winner[pred_col]:.2f}**\nOrt. Gerçek: **%{winner[act_col]:.2f}**\nOrt. Sapma: **{winner[err_col]:.2f} {unit}**")

                show_champion_card(c_best1, "PPK Faizi", "err_ppk", "Puan", "🏦", "tahmin_ppk_faiz", "PPK Faizi")
                show_champion_card(c_best2, "Aylık Enflasyon", "err_enf_ay", "Puan", "📉", "tahmin_aylik_enf", "Aylık TÜFE")
                show_champion_card(c_best3, "Yıllık Enflasyon", "err_enf_yil", "Puan", "🏷️", "val_enf_yil", "Yıllık TÜFE")
                st.caption(f"*Analiz {p_start} ile {p_end} arasındaki dönemleri kapsar.*")
            else: st.info("Seçilen tarih aralığında gerçekleşmiş veri bulunamadı.")
        else: st.warning("Gerçekleşen piyasa verileri çekilemediği için performans analizi yapılamıyor.")
        st.markdown("---")

        with st.sidebar:
            st.markdown("### 🔍 Dashboard Filtreleri")
            x_axis_mode = st.radio("Grafik Görünümü (X Ekseni)", ["📅 Hedef Dönem (Vade)", "⏳ Tahmin Tarihi (Revizyon)"])
            st.markdown("---")
            calc_method = st.radio("Medyan Hesaplama", ["Otomatik", "Manuel"])
            manual_median_val = 0.0 if calc_method == "Otomatik" else st.number_input("Manuel Değer", step=0.01, format="%.2f")
            st.markdown("---")
            cat_filter = st.multiselect("Kategori", ["Bireysel", "Kurumsal"], default=["Bireysel", "Kurumsal"])
            df_filt_base = df_latest[df_latest['kategori'].isin(cat_filter)]
            avail_src = sorted(df_filt_base['anket_kaynagi'].astype(str).unique())
            src_filter = st.multiselect("Kaynak", avail_src, default=avail_src)
            df_filt_src = df_filt_base[df_filt_base['anket_kaynagi'].isin(src_filter)]
            avail_usr = sorted(df_filt_src['gorunen_isim'].unique())
            usr_filter = st.multiselect("Katılımcı", avail_usr, default=avail_usr)
            avail_yr = sorted(df_filt_src['yil'].unique())
            yr_filter = st.multiselect("Yıl", avail_yr, default=avail_yr)

        is_single_user = (len(usr_filter) == 1)
        if is_single_user:
            target_df = df_history[df_history['gorunen_isim'].isin(usr_filter) & df_history['yil'].isin(yr_filter)].copy()
            x_axis_col = "tahmin_tarihi"; x_label = "Tahmin Giriş Tarihi"; sort_col = "tahmin_tarihi"; tick_format = "%d-%m-%Y"
        else:
            target_df = df_latest[df_latest['kategori'].isin(cat_filter) & df_latest['anket_kaynagi'].isin(src_filter) & df_latest['gorunen_isim'].isin(usr_filter) & df_latest['yil'].isin(yr_filter)].copy()
            x_axis_col = "donem"; x_label = "Hedef Dönem"; sort_col = "donem_date"; tick_format = None

        if target_df.empty: st.warning("Seçilen filtrelerde veri bulunamadı."); st.stop()

        tabs = st.tabs(["📈 Zaman Serisi", "📍 Dağılım Analizi", "📦 Kutu Grafiği"])
        with tabs[0]:
            def plot(y, min_c, max_c, tit, real_key=None):
                chart_data = target_df.sort_values(sort_col)
                fig = px.line(chart_data, x=x_axis_col, y=y, color="gorunen_isim" if not is_single_user else "donem", markers=True, title=tit, hover_data=["hover_text"])
                if tick_format: fig.update_xaxes(tickformat=tick_format)
                if x_axis_mode.startswith("📅") and real_key and realized_dict:
                    real_df_data = []
                    for d, vals in realized_dict.items():
                        if vals.get(real_key) is not None: real_df_data.append({'donem': d, 'deger': vals[real_key]})
                    if real_df_data:
                        real_df_p = pd.DataFrame(real_df_data).sort_values('donem')
                        min_d = chart_data['donem'].min(); max_d = chart_data['donem'].max()
                        real_df_p = real_df_p[(real_df_p['donem'] >= min_d) & (real_df_p['donem'] <= max_d)]
                        if not real_df_p.empty: fig.add_trace(go.Scatter(x=real_df_p['donem'], y=real_df_p['deger'], mode='lines+markers', name='GERÇEKLEŞEN', line=dict(color='black', width=4, dash='dot'), marker=dict(size=8, color='black', symbol='x')))
                dfr = chart_data.dropna(subset=[min_c, max_c])
                if not dfr.empty:
                    grp = "donem" if is_single_user else "gorunen_isim"
                    for g in dfr[grp].unique():
                        ud = dfr[dfr[grp] == g]
                        fig.add_trace(go.Scatter(x=ud[x_axis_col], y=ud[y], mode='markers', error_y=dict(type='data', symmetric=False, array=ud[max_c]-ud[y], arrayminus=ud[y]-ud[min_c], color='gray', width=1), showlegend=False, hoverinfo='skip', marker=dict(size=0, opacity=0)))
                st.plotly_chart(fig, use_container_width=True)
            
            c1, c2 = st.columns(2)
            with c1:
                plot("tahmin_ppk_faiz", "min_ppk_faiz", "max_ppk_faiz", "PPK Karar", "ppk")
            with c2:
                plot("tahmin_yilsonu_faiz", "min_yilsonu_faiz", "max_yilsonu_faiz", "Sene Sonu Faiz", None)
                
            c3, c4 = st.columns(2)
            with c3:
                plot("tahmin_aylik_enf", "min_aylik_enf", "max_aylik_enf", "Aylık Enf", "enf_ay")
            with c4:
                plot("tahmin_yilsonu_enf", "min_yilsonu_enf", "max_yilsonu_enf", "YS Enf (Veya Yıllık)", "enf_yil")

        with tabs[1]:
            pers = sorted(list(target_df['donem'].unique()), reverse=True)
            if not pers: st.stop()
            tp = st.selectbox("Dönem Seç", pers, key="dp")
            dp = target_df[target_df['donem'] == tp].copy()
            met_map = {"PPK": "tahmin_ppk_faiz", "Ay Enf": "tahmin_aylik_enf", "YS Enf": "tahmin_yilsonu_enf"}
            sm = st.radio("Metrik", list(met_map.keys()), horizontal=True)
            mc = met_map[sm]
            dp = dp.dropna(subset=[mc])
            if len(dp)>0:
                mv = manual_median_val if calc_method == "Manuel" else dp[mc].median()
                dp = dp.sort_values(by=mc)
                fig = go.Figure()
                y_val = dp['tahmin_tarihi'].dt.strftime('%d-%m-%Y') if (is_single_user) else dp['gorunen_isim']
                fig.add_trace(go.Scatter(x=dp[mc], y=y_val, mode='markers', marker=dict(size=14, color='#1976D2', line=dict(width=1, color='white')), name='Tahmin', text=[f"%{v:.2f}" for v in dp[mc]], hoverinfo='text'))
                fig.add_vline(x=mv, line_width=3, line_color="red")
                fig.add_annotation(x=mv, y=-0.1, text=f"MEDYAN %{mv:.2f}", showarrow=False, font=dict(color="red", size=14, weight="bold"), yref="paper")
                if realized_dict and tp in realized_dict:
                     real_key_map = {"PPK": "ppk", "Ay Enf": "enf_ay", "YS Enf": "enf_yil"}
                     rv = realized_dict[tp].get(real_key_map[sm])
                     if rv is not None: fig.add_vline(x=rv, line_width=3, line_color="black", line_dash="dash"); fig.add_annotation(x=rv, y=-0.2, text=f"GERÇEK %{rv:.2f}", showarrow=False, font=dict(color="black", size=12), yref="paper")
                fig.update_layout(title=f"{sm} Dağılım ({tp})", height=max(500, len(dp)*35))
                st.plotly_chart(fig, use_container_width=True)
            else: st.info("Bu metrik için veri yok")

        with tabs[2]:
            mb = {"PPK": "tahmin_ppk_faiz", "Ay Enf": "tahmin_aylik_enf", "YS Enf": "tahmin_yilsonu_enf"}
            sb = st.selectbox("Veri Seti", list(mb.keys()))
            fig = px.box(target_df.sort_values("donem_date"), x="donem", y=mb[sb], color="donem", title=f"{sb} Dağılımı")
            st.plotly_chart(fig, use_container_width=True)

# ========================================================
# SAYFA: ISI HARİTASI
# ========================================================
elif page == "🔥 Isı Haritası":
    st.header("🔥 Tahmin Isı Haritası")
    res_t = supabase.table(TABLE_TAHMIN).select("*").order("tahmin_tarihi", desc=True).limit(2000).execute()
    df_t = pd.DataFrame(res_t.data)
    res_k = supabase.table(TABLE_KATILIMCI).select("ad_soyad", "anket_kaynagi").execute()
    df_k = pd.DataFrame(res_k.data)

    if not df_t.empty and not df_k.empty:
        df_t = clean_and_sort_data(df_t)
        df_t['tahmin_tarihi'] = pd.to_datetime(df_t['tahmin_tarihi'])
        df_t = df_t.sort_values(by='tahmin_tarihi')
        df_full = pd.merge(df_t, df_k, left_on="kullanici_adi", right_on="ad_soyad", how="inner")
        df_full['gorunen_isim'] = df_full.apply(lambda x: f"{x['kullanici_adi']} ({x['anket_kaynagi']})" if pd.notnull(x['anket_kaynagi']) and x['anket_kaynagi'] != '' else x['kullanici_adi'], axis=1)

        with st.expander("⚙️ Harita Ayarları", expanded=True):
            view_mode = st.radio("Görünüm Modu", ["📅 Hedef Dönem Karşılaştırması", "⏳ Zaman İçindeki Değişim (Revizyon)"], horizontal=True)
            st.markdown("---")
            c1, c2, c3 = st.columns(3)
            metrics = {"PPK Faizi": "tahmin_ppk_faiz", "Yıl Sonu Faiz": "tahmin_yilsonu_faiz", "Aylık Enflasyon": "tahmin_aylik_enf", "Yıl Sonu Enflasyon": "tahmin_yilsonu_enf"}
            sel_metric_label = c1.selectbox("Veri Seti", list(metrics.keys()))
            sel_metric = metrics[sel_metric_label]
            
            all_users = sorted(df_full['gorunen_isim'].unique())
            sel_users = c2.multiselect("Katılımcılar", all_users, default=all_users[:10] if len(all_users)>0 else [])
            all_periods = sorted(df_full['donem'].unique(), reverse=True)
            
            if view_mode.startswith("📅"):
                sel_periods = c3.multiselect("Hedef Dönemler", all_periods, default=all_periods[:6] if len(all_periods)>0 else [])
                if not sel_users or not sel_periods: st.stop()
                df_f = df_full[df_full['gorunen_isim'].isin(sel_users) & df_full['donem'].isin(sel_periods)].copy()
                df_f = df_f.sort_values(by='tahmin_tarihi').drop_duplicates(subset=['kullanici_adi', 'donem'], keep='last')
                piv_col = 'donem'
            else:
                target_period = c3.selectbox("Hangi Hedefin Geçmişini İzliceksiniz?", all_periods)
                time_granularity = c3.radio("Zaman Dilimi", ["🗓️ Aylık (Son Veri)", "📆 Günlük (Detaylı)"])
                if not sel_users or not target_period: st.stop()
                df_f = df_full[df_full['gorunen_isim'].isin(sel_users) & (df_full['donem'] == target_period)].copy()
                if "Günlük" in time_granularity: df_f['tahmin_zaman'] = df_f['tahmin_tarihi'].dt.strftime('%Y-%m-%d')
                else: df_f['tahmin_zaman'] = df_f['tahmin_tarihi'].dt.strftime('%Y-%m')
                df_f = df_f.sort_values(by='tahmin_tarihi').drop_duplicates(subset=['kullanici_adi', 'tahmin_zaman'], keep='last')
                piv_col = 'tahmin_zaman'

        if df_f.empty: st.warning("Veri yok."); st.stop()
        pivot_df = df_f.pivot(index='gorunen_isim', columns=piv_col, values=sel_metric)
        pivot_df = pivot_df.reindex(columns=sorted(pivot_df.columns))

        def highlight(data):
            styles = pd.DataFrame('', index=data.index, columns=data.columns)
            for idx, row in data.iterrows():
                prev = None; first = False
                for col in data.columns:
                    val = row[col]
                    if pd.isna(val): continue
                    st = ''
                    if not first: st='background-color: #FFF9C4; color: black; font-weight: bold; border: 1px solid white;'; first=True
                    else:
                        if prev is not None:
                            if val > prev: st='background-color: #FFCDD2; color: #B71C1C; font-weight: bold; border: 1px solid white;'
                            elif val < prev: st='background-color: #C8E6C9; color: #1B5E20; font-weight: bold; border: 1px solid white;'
                            else: st='background-color: #FFF9C4; color: black; font-weight: bold; border: 1px solid white;'
                    styles.at[idx, col] = st
                    prev = val
            return styles

        st.markdown(f"### 🔥 {sel_metric_label} Analizi")
        st.dataframe(pivot_df.style.apply(highlight, axis=None).format("{:.2f}"), use_container_width=True, height=len(sel_users)*50+100)
        st.caption("🟡: İlk Veri / Değişim Yok | 🔴: Yükseliş | 🟢: Düşüş")
    else: st.info("Veri yok.")

# ========================================================
# SAYFA: TOPLU VERİ YÜKLEME (EXCEL)
# ========================================================
elif page == "📥 Toplu Veri Yükleme (Excel)":
    st.header("📥 Toplu Veri Yükleme")
    st.info("Bu alandan çok sayıda veriyi Excel formatında yükleyebilirsiniz. Sistem mevcut kayıtları kontrol eder ve onayınızı ister.")

    def generate_excel_template():
        # Güncellenmiş Şablon (Medyan, Min, Max alanları dahil)
        cols = [
            "Katılımcı Adı", "Dönem (YYYY-AA)", "Tarih (YYYY-AA-GG)", "Kategori", "Link", 
            "PPK Medyan", "PPK Min", "PPK Max", 
            "Yıl Sonu Faiz Medyan", "Yıl Sonu Faiz Min", "Yıl Sonu Faiz Max",
            "Aylık Enf Medyan", "Aylık Enf Min", "Aylık Enf Max",
            "Yıllık Enf Medyan", "Yıllık Enf Min", "Yıllık Enf Max",
            "Yıl Sonu Enf Medyan", "Yıl Sonu Enf Min", "Yıl Sonu Enf Max",
            "N Sayısı",
            "Gelecek Hedef Dönem", "Gelecek Tahmin (Enf)", "Gelecek Tahmin (PPK)"
        ]
        df_temp = pd.DataFrame(columns=cols)
        # Örnek Satır
        df_temp.loc[0] = [
            "Örnek Banka", "2025-10", "2025-10-15", "Kurumsal", "", 
            45.0, 42.0, 48.0, 
            40.0, 38.0, 42.0,
            1.5, 1.2, 1.8,
            30.0, 28.0, 32.0,
            35.0, 33.0, 37.0,
            15,
            "2026-12", 25.0, 35.0
        ]
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            df_temp.to_excel(writer, index=False, sheet_name='Veri_Girisi')
            worksheet = writer.sheets['Veri_Girisi']
            for i, col in enumerate(cols): worksheet.set_column(i, i, 20)
        return output.getvalue()

    c_dl, c_up = st.columns([1, 2])
    with c_dl:
        st.subheader("1. Şablon İndir")
        st.download_button(label="📥 Excel Şablonunu İndir", data=generate_excel_template(), file_name="Veri_Yukleme_Sablonu_v3.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary")
    with c_up:
        st.subheader("2. Dosya Yükle ve Kontrol Et")
        uploaded_file = st.file_uploader("Excel Dosyası Seç (.xlsx)", type=["xlsx"])

    if uploaded_file:
        try:
            df_upload = pd.read_excel(uploaded_file)
            required_cols = ["Katılımcı Adı", "Dönem (YYYY-AA)", "Tarih (YYYY-AA-GG)"]
            if not all(col in df_upload.columns for col in required_cols): st.error("Excel formatı hatalı! Lütfen güncel şablonu indirip tekrar deneyin."); st.stop()
            st.write("📋 **Yüklenecek Veri Önizlemesi:**"); st.dataframe(df_upload.head(3), use_container_width=True)
            
            excel_user_data = {} 
            for _, row in df_upload.iterrows():
                nm = normalize_name(str(row["Katılımcı Adı"]))
                cat = str(row.get("Kategori", "Bireysel"))
                excel_user_data[nm] = cat
            res_k = supabase.table(TABLE_KATILIMCI).select("ad_soyad").execute()
            db_users = set([r['ad_soyad'] for r in res_k.data])
            users_to_add = []
            for nm, cat in excel_user_data.items():
                if nm not in db_users: users_to_add.append({"ad_soyad": nm, "kategori": cat})
            
            if 'check_done' not in st.session_state: st.session_state['check_done'] = False
            
            if st.button("🔍 Veritabanı ile Karşılaştır"):
                with st.spinner("Mevcut kayıtlar kontrol ediliyor..."):
                    res = supabase.table(TABLE_TAHMIN).select("kullanici_adi, donem").execute()
                    existing_set = set()
                    for r in res.data: existing_set.add((r['kullanici_adi'], r['donem']))
                    duplicates = []; new_records = []
                    for index, row in df_upload.iterrows():
                        u_name = str(row["Katılımcı Adı"]).strip(); u_period = str(row["Dönem (YYYY-AA)"]).strip()
                        u_future_period = str(row.get("Gelecek Hedef Dönem", "")).strip()
                        if (u_name, u_period) in existing_set: duplicates.append(f"{u_name} - {u_period}")
                        else: new_records.append(f"{u_name} - {u_period}")
                        if u_future_period and u_future_period.lower() != "nan" and u_future_period != "":
                             if (u_name, u_future_period) in existing_set: duplicates.append(f"{u_name} - {u_future_period} (Gelecek Tahmin)")
                             else: new_records.append(f"{u_name} - {u_future_period} (Gelecek Tahmin)")
                    st.session_state['duplicates'] = duplicates; st.session_state['new_count'] = len(new_records); st.session_state['check_done'] = True

            if st.session_state.get('check_done'):
                dups = st.session_state['duplicates']; cnt_new = st.session_state['new_count']; cnt_dup = len(dups)
                if users_to_add:
                    st.info(f"🆕 **{len(users_to_add)}** Yeni Katılımcı tespit edildi. İşlem sırasında otomatik eklenecekler.")
                    with st.expander("Yeni Eklenecek Katılımcılar"): st.write([u['ad_soyad'] for u in users_to_add])
                else: st.success("✅ Tüm katılımcılar sistemde kayıtlı.")

                st.markdown("---"); c1, c2 = st.columns(2)
                c1.info(f"🆕 **{cnt_new}** adet yeni kayıt oluşturulacak.")
                confirm_overwrite = True 
                if cnt_dup > 0:
                    c2.warning(f"⚠️ **{cnt_dup}** adet kayıt veritabanında ZATEN MEVCUT!")
                    with st.expander("Çakışan Kayıtları Gör"): st.write(dups)
                    confirm_overwrite = st.checkbox("Mevcut kayıtların üzerine yazılmasını onaylıyorum ✅", value=False)
                    if not confirm_overwrite: st.error("Devam etmek için yukarıdaki onay kutusunu işaretleyin.")
                
                if st.button("🚀 İşlemi Başlat", type="primary", disabled=(cnt_dup > 0 and not confirm_overwrite)):
                    progress_bar = st.progress(0); success_count = 0; total_rows = len(df_upload)
                    if users_to_add:
                        for new_u in users_to_add:
                            try: supabase.table(TABLE_KATILIMCI).insert(new_u).execute()
                            except: pass
                    for index, row in df_upload.iterrows():
                        try:
                            user = str(row["Katılımcı Adı"]).strip(); period = str(row["Dönem (YYYY-AA)"]).strip(); cat = str(row.get("Kategori", "Bireysel")).strip()
                            link = str(row["Link"]) if pd.notnull(row["Link"]) else None
                            raw_date = row["Tarih (YYYY-AA-GG)"]
                            if isinstance(raw_date, pd.Timestamp): forecast_date = raw_date.strftime("%Y-%m-%d")
                            else: forecast_date = str(raw_date).split()[0]
                            def cv(val): 
                                try: v = float(val); return v if pd.notnull(v) else None
                                except: return None
                            
                            data_main = {
                                "tahmin_ppk_faiz": cv(row.get("PPK Medyan")),
                                "min_ppk_faiz": cv(row.get("PPK Min")),
                                "max_ppk_faiz": cv(row.get("PPK Max")),
                                "tahmin_yilsonu_faiz": cv(row.get("Yıl Sonu Faiz Medyan")),
                                "min_yilsonu_faiz": cv(row.get("Yıl Sonu Faiz Min")),
                                "max_yilsonu_faiz": cv(row.get("Yıl Sonu Faiz Max")),
                                "tahmin_aylik_enf": cv(row.get("Aylık Enf Medyan")),
                                "min_aylik_enf": cv(row.get("Aylık Enf Min")),
                                "max_aylik_enf": cv(row.get("Aylık Enf Max")),
                                "tahmin_yillik_enf": cv(row.get("Yıllık Enf Medyan")),
                                "min_yillik_enf": cv(row.get("Yıllık Enf Min")),
                                "max_yillik_enf": cv(row.get("Yıllık Enf Max")),
                                "tahmin_yilsonu_enf": cv(row.get("Yıl Sonu Enf Medyan")),
                                "min_yilsonu_enf": cv(row.get("Yıl Sonu Enf Min")),
                                "max_yilsonu_enf": cv(row.get("Yıl Sonu Enf Max")),
                                "katilimci_sayisi": int(cv(row.get("N Sayısı")) or 0)
                            }
                            upsert_tahmin(user, period, cat, forecast_date, link, data_main)
                            success_count += 1
                            
                            fut_period = str(row.get("Gelecek Hedef Dönem", "")).strip()
                            if fut_period and fut_period.lower() != "nan" and fut_period != "":
                                fut_enf = cv(row.get("Gelecek Tahmin (Enf)")); fut_ppk = cv(row.get("Gelecek Tahmin (PPK)"))
                                if fut_enf or fut_ppk:
                                    data_future = {"katilimci_sayisi": int(cv(row.get("N Sayısı")) or 0)}
                                    if fut_enf: data_future["tahmin_yilsonu_enf"] = fut_enf; data_future["tahmin_yillik_enf"] = fut_enf
                                    if fut_ppk: data_future["tahmin_ppk_faiz"] = fut_ppk; data_future["tahmin_yilsonu_faiz"] = fut_ppk
                                    upsert_tahmin(user, fut_period, cat, forecast_date, link, data_future)
                        except Exception as e: st.error(f"Satır {index+1} Hatası: {e}")
                        progress_bar.progress((index + 1) / total_rows)
                    st.success(f"İşlem Tamamlandı! {success_count} satır işlendi."); time.sleep(2); del st.session_state['check_done']; del st.session_state['duplicates']; st.rerun()
        except Exception as e: st.error(f"Dosya okuma hatası: {e}")

# ========================================================
# SAYFA: VERİ GİRİŞ EKRANLARI
# ========================================================
elif page in ["PPK Girişi", "Enflasyon Girişi"]:
    st.header(f"➕ {page}")
    with st.container():
        with st.form("entry_form"):
            st.subheader("1. Ana Tahmin"); c1, c2, c3 = st.columns([2, 1, 1])
            with c1: user, cat, disp = get_participant_selection()
            def_idx = tum_donemler.index("2025-01") if "2025-01" in tum_donemler else 0
            with c2: donem = st.selectbox("Dönem (Cari)", tum_donemler, index=def_idx)
            with c3: tarih = st.date_input("Giriş Tarihi", datetime.date.today())
            link = st.text_input("Kaynak Linki (Opsiyonel)")
            st.markdown("---"); data = {}; kat_sayisi = 0
            if page == "PPK Girişi":
                c1, c2 = st.columns(2); r1 = c1.text_input("Aralık (42-45)", key="r1"); v1 = c1.number_input("Medyan %", step=0.25); r2 = c2.text_input("Aralık YS", key="r2"); v2 = c2.number_input("YS Medyan %", step=0.25)
                with st.expander("Detaylar (Min/Max/N)"):
                    ec1, ec2, ec3 = st.columns(3); mn1 = ec1.number_input("Min", step=0.25); mx1 = ec1.number_input("Max", step=0.25); mn2 = ec2.number_input("Min YS", step=0.25); mx2 = ec2.number_input("Max YS", step=0.25); kat_sayisi = ec3.number_input("N", step=1)
                md, mn, mx, ok = parse_range_input(r1, v1); 
                if ok: v1, mn1, mx1 = md, mn, mx
                md2, mn2, mx2, ok2 = parse_range_input(r2, v2)
                if ok2: v2, mn2, mx2 = md2, mn2, mx2
                data = {"tahmin_ppk_faiz": v1, "min_ppk_faiz": mn1, "max_ppk_faiz": mx1, "tahmin_yilsonu_faiz": v2, "min_yilsonu_faiz": mn2, "max_yilsonu_faiz": mx2}
            else: # Enflasyon Girişi
                c1, c2, c3 = st.columns(3); r1 = c1.text_input("Aralık Ay", key="r1"); v1 = c1.number_input("Ay Medyan", step=0.01, format="%.2f"); r2 = c2.text_input("Aralık Yıllık", key="r2"); v2 = c2.number_input("Yıllık Medyan", step=0.01, format="%.2f"); r3 = c3.text_input("Aralık YS", key="r3"); v3 = c3.number_input("YS Medyan", step=0.01, format="%.2f")
                with st.expander("Detaylar (Min/Max/N)"):
                    ec1, ec2, ec3 = st.columns(3); mn1 = ec1.number_input("Min Ay", step=0.01); mx1 = ec1.number_input("Max Ay", step=0.01); mn2 = ec2.number_input("Min Yıl", step=0.01); mx2 = ec2.number_input("Max Yıl", step=0.01); mn3 = ec3.number_input("Min YS", step=0.01); mx3 = ec3.number_input("Max YS", step=0.01); kat_sayisi = st.number_input("N", step=1)
                md1, mn1, mx1, ok1 = parse_range_input(r1, v1); 
                if ok1: v1, mn1, mx1 = md1, mn1, mx1
                md2, mn2, mx2, ok2 = parse_range_input(r2, v2)
                if ok2: v2, mn2, mx2 = md2, mn2, mx2
                md3, mn3, mx3, ok3 = parse_range_input(r3, v3)
                if ok3: v3, mn3, mx3 = md3, mn3, mx3
                data = {"tahmin_aylik_enf": v1, "min_aylik_enf": mn1, "max_aylik_enf": mx1, "tahmin_yillik_enf": v2, "min_yillik_enf": mn2, "max_yillik_enf": mx2, "tahmin_yilsonu_enf": v3, "min_yilsonu_enf": mn3, "max_yilsonu_enf": mx3}
            data["katilimci_sayisi"] = int(kat_sayisi) if kat_sayisi > 0 else 0
            
            extra_future_data = None; future_donem = None
            st.markdown("---"); st.markdown("#### 📅 İleri Vadeli Beklenti (Opsiyonel)"); st.caption("Örn: Rapor şu anki ayı (Ekim) ele alıyor ama 'Gelecek Yıl Sonu' için de bir tahmin içeriyor.")
            fe1, fe2 = st.columns(2)
            try: curr_year = int(donem.split('-')[0]); next_december = f"{curr_year + 1}-12"; f_idx = tum_donemler.index(next_december) if next_december in tum_donemler else 0
            except: f_idx = 0
            future_donem = fe1.selectbox("Hedef Dönem (Gelecek)", tum_donemler, index=f_idx)
            if page == "PPK Girişi":
                future_val = fe2.number_input("Gelecek Dönem Politika Faizi Beklentisi (%)", step=0.25, format="%.2f")
                if future_val > 0: extra_future_data = {"tahmin_ppk_faiz": future_val, "tahmin_yilsonu_faiz": future_val, "katilimci_sayisi": int(kat_sayisi)}
            else: # Enflasyon
                future_val = fe2.number_input("Gelecek Dönem Enflasyon Beklentisi (%)", step=0.01, format="%.2f")
                if future_val > 0: extra_future_data = {"tahmin_yilsonu_enf": future_val, "tahmin_yillik_enf": future_val, "katilimci_sayisi": int(kat_sayisi)}

            if st.form_submit_button("✅ Kaydet"):
                if user:
                    upsert_tahmin(user, donem, cat, tarih, link, data)
                    if extra_future_data and future_donem and future_val > 0:
                        upsert_tahmin(user, future_donem, cat, tarih, link, extra_future_data)
                        st.toast(f"Kaydedildi! (Ana Dönem: {donem} + İleri Dönem: {future_donem})", icon="🎉")
                    else: st.toast(f"Kaydedildi! ({donem})", icon="🎉")
                else: st.error("Kullanıcı Seçiniz")

# ========================================================
# SAYFA: KATILIMCI YÖNETİMİ
# ========================================================
elif page == "Katılımcı Yönetimi":
    st.header("👥 Katılımcı Yönetimi")
    
    # 1. Yeni Kişi Ekleme
    with st.expander("➕ Yeni Kişi Ekle", expanded=False):
        with st.form("new_kat"):
            c1, c2 = st.columns(2); ad = c1.text_input("Ad / Kurum Adı"); cat = c2.radio("Kategori", ["Bireysel", "Kurumsal"], horizontal=True); src = st.text_input("Kaynak (Kurum ise Banka Adı vb.)")
            if st.form_submit_button("Ekle"):
                if ad:
                    try: 
                        clean_name = normalize_name(ad)
                        check = supabase.table(TABLE_KATILIMCI).select("*").eq("ad_soyad", clean_name).execute()
                        if check.data: st.error("Bu isimde bir katılımcı zaten var!")
                        else: supabase.table(TABLE_KATILIMCI).insert({"ad_soyad": clean_name, "kategori": cat, "anket_kaynagi": src or None}).execute(); st.success(f"{clean_name} başarıyla eklendi!"); time.sleep(1); st.rerun()
                    except Exception as e: st.error(f"Hata: {e}")
                else: st.warning("Lütfen bir isim giriniz.")

    # 2. Katılımcı Listesi ve Silme
    st.markdown("---"); res = supabase.table(TABLE_KATILIMCI).select("*").order("ad_soyad").execute(); df_kat = pd.DataFrame(res.data)
    if not df_kat.empty:
        c_list, c_action = st.columns([2, 1])
        with c_list: st.subheader("Mevcut Katılımcılar"); st.dataframe(df_kat, use_container_width=True)
        with c_action:
            st.subheader("🚫 Kişi Silme"); ks = st.selectbox("Silinecek Kişiyi Seç", df_kat["ad_soyad"].unique()); st.warning(f"Dikkat: '{ks}' silindiğinde, ona ait geçmiş tüm tahmin verileri de silinecektir.")
            if st.button("🗑️ Kişiyi ve Verilerini Sil", type="primary"):
                supabase.table(TABLE_TAHMIN).delete().eq("kullanici_adi", ks).execute()
                supabase.table(TABLE_KATILIMCI).delete().eq("ad_soyad", ks).execute()
                st.toast(f"{ks} ve tüm verileri silindi.", icon="👋"); time.sleep(1); st.rerun()

    # 3. Orphan Data Temizliği
    st.markdown("---"); st.subheader("🛠️ Veritabanı Bakımı")
    with st.expander("🧹 Yetim Verileri (Orphan Data) Temizle"):
        st.info("Bu işlem, 'Katılımcılar' listesinde kaydı olmayan ancak 'Tahminler' tablosunda verisi kalmış (silinmiş kullanıcıların artığı) kayıtları tespit eder ve temizler.")
        if st.button("Taramayı Başlat"):
            with st.spinner("Veritabanı taranıyor..."):
                valid_users = set(df_kat["ad_soyad"].unique()) if not df_kat.empty else set()
                res_t = supabase.table(TABLE_TAHMIN).select("kullanici_adi").execute(); df_t_users = pd.DataFrame(res_t.data)
                if not df_t_users.empty:
                    existing_users = set(df_t_users["kullanici_adi"].unique())
                    orphans = existing_users - valid_users
                    if orphans:
                        st.error(f"Tespit edilen yetim kullanıcı verileri: {len(orphans)} adet"); st.write(list(orphans))
                        if st.button("🗑️ Hepsini Kalıcı Olarak Sil"):
                            for orphan_user in orphans: supabase.table(TABLE_TAHMIN).delete().eq("kullanici_adi", orphan_user).execute()
                            st.success("Temizlik tamamlandı! Veritabanı tutarlı."); time.sleep(2); st.rerun()
                    else: st.success("✅ Veritabanı temiz. Yetim veri bulunamadı.")
                else: st.info("Tahmin tablosu boş.")
    
    # 4. TEHLİKE BÖLGESİ (E-POSTA ONAYLI SİLME)
    st.markdown("---")
    with st.expander("🚨 TEHLİKE BÖLGESİ: TÜM VERİLERİ SİL"):
        st.error("Bu işlem veritabanındaki TÜM tahmin verilerini ve katılımcıları kalıcı olarak siler.")
        
        if 'delete_stage' not in st.session_state: st.session_state['delete_stage'] = 0 
        
        if st.session_state['delete_stage'] == 0:
            if st.button("🔥 Sıfırlama Talebi Oluştur"):
                st.session_state['delete_stage'] = 1
                st.rerun()
        
        if st.session_state['delete_stage'] == 1:
            st.info("İşleme devam etmek için Yönetici Şifresini giriniz.")
            admin_pwd = st.text_input("Yönetici Şifresi", type="password")
            
            if st.button("Onay Kodu Gönder"):
                if admin_pwd == SITE_SIFRESI:
                    otp_code = str(random.randint(100000, 999999))
                    st.session_state['generated_otp'] = otp_code
                    success, msg = send_verification_email(otp_code)
                    if success:
                        st.success(f"Onay kodu s.idrisoglu@gmail.com adresine gönderildi.")
                        st.session_state['delete_stage'] = 2
                        st.rerun()
                    else:
                        st.error(f"E-posta gönderilemedi: {msg}")
                else:
                    st.error("Hatalı Şifre!")
            if st.button("İptal"):
                st.session_state['delete_stage'] = 0
                st.rerun()

        if st.session_state['delete_stage'] == 2:
            st.warning("Lütfen e-postanıza gelen 6 haneli kodu giriniz.")
            user_otp = st.text_input("Onay Kodu")
            
            if st.button("🔥 ONKAYLA VE SİL 🔥", type="primary"):
                if user_otp == st.session_state.get('generated_otp'):
                    supabase.table(TABLE_TAHMIN).delete().neq("id", 0).execute()
                    supabase.table(TABLE_KATILIMCI).delete().neq("id", 0).execute()
                    st.success("Tüm veriler başarıyla silindi ve sistem sıfırlandı.")
                    st.session_state['delete_stage'] = 0
                    time.sleep(3)
                    st.rerun()
                else:
                    st.error("Hatalı onay kodu!")
            if st.button("Vazgeç"):
                st.session_state['delete_stage'] = 0
                st.rerun()
